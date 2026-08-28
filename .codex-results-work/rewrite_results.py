from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Gent-Terneuzen canal\.codex-results-work\6. Results & Discussion.docx")
OUTPUT = Path(r"C:\Gent-Terneuzen canal\.codex-results-work\6. Results & Discussion - rewritten.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=75, bottom=55, end=75):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.4):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths, numeric_from=1, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table.rows[0].height = Pt(23)
    repeat_table_header(table.rows[0])

    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        set_cell_shading(cell, "D9E2F3")
        set_cell_text(
            cell,
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT if i < numeric_from else WD_ALIGN_PARAGRAPH.CENTER,
            size=font_size,
        )

    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, (value, width) in enumerate(zip(row_values, widths)):
            cell = row.cells[i]
            cell.width = Inches(width)
            set_cell_text(
                cell,
                value,
                align=WD_ALIGN_PARAGRAPH.LEFT if i < numeric_from else WD_ALIGN_PARAGRAPH.CENTER,
                size=font_size,
            )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 0.2
    return table


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.widow_control = True
    p.add_run(text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.font.size = Pt(9)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    return p


doc = Document(SOURCE)

# Remove the superseded body, including tracked insertions/deletions, while retaining
# the source section geometry, footer, theme, and style definitions.
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

# Deliver a clean manuscript section rather than carrying the source's tracked-revision state.
settings = doc.settings._element
for tag in ("trackRevisions", "doNotTrackMoves", "doNotTrackFormatting"):
    node = settings.find(qn(f"w:{tag}"))
    if node is not None:
        settings.remove(node)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.08
normal.paragraph_format.space_after = Pt(6)

for name, size, before, after in (
    ("Heading 1", 18, 0, 8),
    ("Heading 2", 13.5, 10, 4),
    ("Heading 3", 11.5, 7, 2),
):
    style = doc.styles[name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

add_heading(doc, "6. Results and Discussion", 1)

add_body(
    doc,
    "All models were evaluated under the same leakage-safe temporal split: January–July 2025 was used for training, August for model selection and, where applicable, residual calibration, and September for final testing. The common test set contained 2,593 forecast origins. Each origin generated a 72-hour trajectory at 15-minute resolution (288 lead steps), corresponding to 746,784 forecast–observation pairs. Because adjacent origins have strongly overlapping histories and target windows, these pairs are repeated forecast evaluations rather than independent statistical samples. Persistence, defined as carrying the latest observed target value through the full horizon, served as the operational reference.",
)
add_body(
    doc,
    "The unrestricted representation supplied up to 72 hours of recent observations, whereas the manually specified representation concentrated the inputs around the physically motivated 42-hour travel-time hypothesis. For classical models, this reduced 1,440 flattened inputs to four scalar features; for the sequence networks and Chronos-2, it changed a 288 × 5 history into a 168 × 4 history. Consequently, the comparison tests two complete representations that differ in dimensionality, history length, and station content. It does not isolate the causal contribution of a single value at exactly 42 hours.",
)

add_heading(doc, "6.1 Overall Matched-Window Performance", 2)
add_body(
    doc,
    "Table 1 summarizes full-horizon September performance. Zero-shot Chronos-2 with the unrestricted 72-hour representation achieved the lowest RMSE (0.2595 mS/cm), the highest R² (0.1631), and the largest reduction in RMSE relative to persistence (22.93%). The manually represented raw TCN was the strongest locally trained model and the second-best configuration overall, with an RMSE of 0.2737 mS/cm and 18.69% skill. Among the classical models, manually represented XGBoost performed best (RMSE 0.2982 mS/cm; 11.43% skill).",
)
add_body(
    doc,
    "Every reported configuration improved RMSE over persistence, although several had negative R². These statements are compatible because the two statistics use different reference predictions: skill compares each forecast with origin-specific persistence, whereas R² compares squared error with a single global-mean predictor. Positive operational skill can therefore coexist with negative R² over the combined set of trajectories.",
)

add_caption(doc, "Table 1. Overall performance across all 72-hour September test trajectories.")
table1_rows = [
    ("Persistence", "—", "Reference", "0.3366", "—", "0.00"),
    ("Ridge", "Unrestricted 72 h", "Aug.-calibrated", "0.3322", "−0.3722", "1.31"),
    ("Ridge", "Manual 42 h", "Aug.-calibrated", "0.3214", "−0.2846", "4.51"),
    ("XGBoost", "Unrestricted 72 h", "Aug.-calibrated", "0.3177", "−0.2548", "5.63"),
    ("XGBoost", "Manual 42 h", "Aug.-calibrated", "0.2982", "−0.1053", "11.43"),
    ("Linear SVR", "Unrestricted 72 h", "Aug.-calibrated", "0.3285", "−0.3417", "2.42"),
    ("Linear SVR", "Manual 42 h", "Aug.-calibrated", "0.3230", "−0.2968", "4.06"),
    ("LSTM", "Unrestricted 72 h", "Raw", "0.3055", "−0.1607", "9.24"),
    ("LSTM", "Manual 42 h", "Raw", "0.2838", "−0.0016", "15.69"),
    ("TCN", "Unrestricted 72 h", "Raw", "0.2882", "−0.0325", "14.40"),
    ("TCN", "Manual 42 h", "Raw", "0.2737", "0.0684", "18.69"),
    ("Chronos-2", "Unrestricted 72 h", "Zero-shot", "0.2595", "0.1631", "22.93"),
    ("Chronos-2", "Unrestricted 72 h", "Aug.-calibrated", "0.2673", "0.1120", "20.61"),
    ("Chronos-2", "Manual 42 h", "Zero-shot", "0.2678", "0.1083", "20.44"),
    ("Chronos-2", "Manual 42 h", "Aug.-calibrated", "0.2695", "0.0969", "19.94"),
]
add_table(
    doc,
    ("Model", "Representation", "Output", "RMSE\n(mS/cm)", "R²", "RMSE skill\n(%)"),
    table1_rows,
    (1.05, 1.35, 1.15, 1.05, 0.75, 1.15),
    numeric_from=3,
    font_size=7.8,
)
add_note(
    doc,
    "Note: RMSE skill = 100 × (1 − model RMSE / persistence RMSE). Classical notebook outputs are the August-calibrated predictions. Raw neural rows are shown because calibration was effectively the identity for both LSTMs and the unrestricted TCN; for the manual TCN, calibration changed RMSE only from 0.2737 to 0.2740 mS/cm. Chronos-2 reports all four prespecified combinations of representation and output processing.",
)

add_heading(doc, "6.2 Lead-Time Behavior", 2)
add_body(
    doc,
    "The relative ranking changed with lead time (Table 2). At 4 and 8 hours, the August-calibrated manual Chronos-2 configuration had the highest checkpoint skill among the Chronos variants (12.41% and 14.25%, respectively), while the manual TCN was strongest among the locally trained models (14.48% and 15.16%). Beyond 8 hours, zero-shot Chronos-2 with unrestricted history improved rapidly relative to persistence: its skill increased from 8.37% at 4 hours to 24.46% at 48 hours. Manual TCN also remained strong, reaching 19.15% at 24 hours and 18.40% at 48 hours.",
)
add_body(
    doc,
    "The classical models showed smaller changes across the selected checkpoints. Manual XGBoost maintained approximately 9–12% skill from 4 to 48 hours, whereas manual Ridge and Linear SVR remained near 4–6%. Thus, the widening advantage of the strongest sequence and foundation-model forecasts was concentrated mainly in the middle and later portions of the 72-hour trajectory rather than at the earliest leads.",
)

add_caption(doc, "Table 2. RMSE skill relative to persistence at selected lead times (%; higher is better).")
table2_rows = [
    ("Ridge", "Manual 42 h", "5.37", "6.03", "5.39", "4.89", "5.31"),
    ("XGBoost", "Manual 42 h", "9.37", "11.39", "11.33", "11.43", "11.61"),
    ("Linear SVR", "Manual 42 h", "5.00", "5.79", "5.07", "4.45", "5.02"),
    ("LSTM (raw)", "Manual 42 h", "13.38", "13.86", "15.20", "15.40", "15.29"),
    ("TCN (raw)", "Manual 42 h", "14.48", "15.16", "17.89", "19.15", "18.40"),
    ("Chronos-2 zero-shot", "Unrestricted 72 h", "8.37", "10.04", "19.58", "21.49", "24.46"),
    ("Chronos-2 calibrated", "Unrestricted 72 h", "9.20", "10.68", "16.57", "17.56", "22.12"),
    ("Chronos-2 zero-shot", "Manual 42 h", "11.87", "14.01", "18.46", "18.08", "23.03"),
    ("Chronos-2 calibrated", "Manual 42 h", "12.41", "14.25", "17.79", "17.71", "22.68"),
]
add_table(
    doc,
    ("Model/output", "Representation", "4 h", "8 h", "16 h", "24 h", "48 h"),
    table2_rows,
    (1.6, 1.3, 0.72, 0.72, 0.72, 0.72, 0.72),
    numeric_from=2,
    font_size=8.0,
)
add_note(doc, "Note: Checkpoint skill is computed from all forecast values up to and including the stated lead time, rather than from a single endpoint.")

add_heading(doc, "6.3 Model-Specific Findings", 2)

add_heading(doc, "6.3.1 Ridge Regression", 3)
add_body(
    doc,
    "Ridge regression provided a regularized linear benchmark. On August validation data, the manual representation reduced RMSE from 0.2934 to 0.2856 mS/cm. This advantage transferred to September: RMSE decreased from 0.3322 to 0.3214 mS/cm and skill increased from 1.31% to 4.51%. The unrestricted model was slightly worse than persistence at the 4-hour checkpoint (−3.15% skill) before becoming positive at longer checkpoints, whereas the manual model remained positive throughout. The result indicates that compact lag-informed inputs improved this linear model, but its remaining error was substantially higher than that of the strongest nonlinear and sequence-based configurations.",
)

add_heading(doc, "6.3.2 XGBoost", 3)
add_body(
    doc,
    "XGBoost was the strongest classical model. The manual representation improved August validation RMSE from 0.3010 to 0.2854 mS/cm and September RMSE from 0.3177 to 0.2982 mS/cm. The September change corresponds to a 6.1% reduction relative to unrestricted XGBoost and a 5.80-percentage-point increase in persistence skill. Its approximately stable skill across the reported checkpoints suggests that the compact nonlinear mapping contributed across much of the horizon rather than only near the nominal 42-hour lag.",
)

add_heading(doc, "6.3.3 Linear Support Vector Regression", 3)
add_body(
    doc,
    "The linear SVR showed the same directional representation effect but a smaller test improvement. Manual inputs reduced August validation RMSE from 0.3042 to 0.2853 mS/cm and September RMSE from 0.3285 to 0.3230 mS/cm. Both versions exceeded persistence overall, with skills of 2.42% and 4.06%, respectively, but neither approached manual XGBoost. The similar September performance of Ridge and linear SVR indicates that changing the linear loss and margin formulation alone did not capture the nonlinear structure exploited by the better-performing models.",
)

add_heading(doc, "6.3.4 Long Short-Term Memory Network", 3)
add_body(
    doc,
    "Validation selected the 96 × 48 manual LSTM (validation RMSE 0.2602 mS/cm) and the 128 × 64 unrestricted LSTM (0.2707 mS/cm) within their respective candidate sets. On September data, the manual LSTM retained the advantage, reducing RMSE from 0.3055 to 0.2838 mS/cm and increasing skill from 9.24% to 15.69%. The fitted residual multipliers were effectively one, so the raw and calibrated LSTM forecasts were numerically equivalent. Validation-based LSTM–TCN blending also assigned the LSTM a weight of 1.00 for both representations; consequently, the blend added no distinct forecast and is not duplicated in the tables.",
)

add_heading(doc, "6.3.5 Temporal Convolutional Network", 3)
add_body(
    doc,
    "Validation selected the manual TCN with 48 channels and kernel size 3 (RMSE 0.2651 mS/cm) and the unrestricted TCN with 32 channels and kernel size 5 (0.2831 mS/cm). The manual TCN achieved September RMSE of 0.2737 mS/cm, compared with 0.2882 mS/cm for the unrestricted version, and was the strongest model trained solely on the study data. Its positive R² (0.0684) further distinguished it from the LSTM and classical models. August calibration left the unrestricted TCN unchanged and slightly increased manual TCN test RMSE from 0.2737 to 0.2740 mS/cm, so there was no test-period gain from this post-processing step.",
)

add_heading(doc, "6.3.6 Chronos-2", 3)
add_body(
    doc,
    "Chronos-2 was evaluated along two independent axes: input representation and output processing. The unrestricted 72-hour and manual 42-hour representations were each run in their native zero-shot form and with lead-specific residual multipliers fitted on August. Thus, zero-shot is not synonymous with unrestricted, and calibrated is not synonymous with manual. The unrestricted zero-shot configuration was the best overall September forecast (RMSE 0.2595 mS/cm; 22.93% skill). The manual zero-shot configuration followed at 0.2678 mS/cm and 20.44% skill.",
)
add_body(
    doc,
    "Calibration favored the manual representation on August (calibrated validation RMSE 0.2744 versus 0.2803 mS/cm for unrestricted inputs), but this ordering did not carry over unchanged to September. Calibration increased full-horizon September RMSE from 0.2595 to 0.2673 mS/cm for unrestricted inputs and from 0.2678 to 0.2695 mS/cm for manual inputs. It improved the 4- and 8-hour checkpoint skill for both representations, but reduced skill at most later checkpoints. The result is consistent with month-specific differences in the optimal correction magnitude and does not support treating the August multipliers as universally transferable.",
)

add_caption(doc, "Table 3. Chronos-2 residual-interval diagnostics on the September test set.")
table3_rows = [
    ("Unrestricted 72 h", "80", "69.82", "0.4806", "0.9719", "69.19"),
    ("Manual 42 h", "80", "69.74", "0.4631", "0.9215", "68.45"),
]
add_table(
    doc,
    ("Representation", "Nominal\ncoverage (%)", "Empirical\ncoverage (%)", "Mean width\n(mS/cm)", "Mean interval\nscore", "72 h endpoint\ncoverage (%)"),
    table3_rows,
    (1.2, 1.0, 1.0, 1.05, 1.1, 1.25),
    numeric_from=1,
    font_size=7.8,
)
add_note(
    doc,
    "Note: Intervals were constructed from August residual quantiles and evaluated on September. Lower interval score is better, conditional on calibration.",
)
add_body(
    doc,
    "The nominal 80% Chronos-2 residual intervals covered only about 69.8% of all September observations for both representations (Table 3). Manual inputs produced slightly narrower intervals and a lower mean interval score, but coverage remained approximately ten percentage points below nominal. The intervals therefore quantify useful spread around the point forecasts but should not be interpreted as calibrated 80% operational uncertainty bands without further recalibration and external testing.",
)

add_heading(doc, "6.4 Interpretation of the 42-Hour Hypothesis", 2)
add_body(
    doc,
    "The manual 42-hour representation improved every classical model and both locally trained sequence architectures. The effect was largest for XGBoost, LSTM, and TCN, and it was reproduced in both August validation and September testing. This consistency supports the practical value of a compact lag-informed representation for models trained on the available study data.",
)
add_body(
    doc,
    "Chronos-2 qualified that pattern. Manual inputs improved its earliest checkpoint skill, but unrestricted zero-shot inputs gave the best full-horizon result and the highest skill from 16 hours onward. The 42-hour hypothesis should therefore be interpreted as a useful feature-engineering prior rather than as evidence that exactly 42 hours is universally optimal. Because the ablation simultaneously changed history length, input dimensionality, and station composition, a causal travel-time claim would require controlled alternatives that vary the lag while holding the remaining representation fixed.",
)

add_heading(doc, "6.5 Limitations and Operational Implications", 2)
add_body(
    doc,
    "The evaluation covers a single held-out month. Although 746,784 forecast–observation pairs were scored, forecasts from adjacent 15-minute origins share most of their history and many target timestamps; this dependence precludes treating individual pairs or origins as independent replicates for conventional significance testing. Rankings may also depend on season, hydrological regime, missing-data patterns, or changes in station behavior. External drivers such as discharge, tides, rainfall, freshwater inflow, temperature, vessel movements, and lock operations were not included and may explain events that history-only models smooth or miss.",
)
add_body(
    doc,
    "The August-to-September degradation of the Chronos-2 calibration factors and the undercoverage of its residual intervals show that point and uncertainty calibration require monitoring under temporal distribution shift. Future evaluation should use multiple blocked months or hydrological events, report day- or event-level uncertainty in performance differences, and recalibrate intervals on rolling or regime-matched windows. Controlled lag ablations are also needed before attaching a unique physical interpretation to the 42-hour setting.",
)
add_body(
    doc,
    "For the present test month, unrestricted zero-shot Chronos-2 provides the strongest full-horizon accuracy, while the manual TCN offers the strongest locally trained alternative. At the earliest checkpoints, manually represented Chronos-2 is comparatively advantageous. These complementary patterns suggest that operational selection should depend on required lead time, computational constraints, calibration stability, and performance on additional independent periods; the current experiment alone is not sufficient for a definitive deployment choice.",
)

# Keep the source's single-section footer and geometry; make sure no accidental
# extra blank paragraphs trail the body.
doc.save(OUTPUT)
print(OUTPUT)
