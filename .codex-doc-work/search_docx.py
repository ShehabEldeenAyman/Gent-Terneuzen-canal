import sys
from pathlib import Path

from docx import Document


def iter_blocks(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        yield f"P{i:04d}", paragraph.style.name, " ".join(paragraph.text.split())
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            text = " || ".join(" ".join(cell.text.split()) for cell in row.cells)
            yield f"T{ti:03d}R{ri:03d}", "Table row", text


def main() -> None:
    path = Path(sys.argv[1])
    terms = [term.casefold() for term in sys.argv[2:]]
    doc = Document(path)
    blocks = list(iter_blocks(doc))
    print(f"FILE: {path}")
    print("HEADINGS:")
    for label, style, text in blocks:
        if text and style.casefold().startswith("heading"):
            print(f"{label}\t{style}\t{text}")
    for term in terms:
        print(f"\nTERM: {term}")
        matches = [i for i, (_, _, text) in enumerate(blocks) if term in text.casefold()]
        if not matches:
            print("NO MATCH")
            continue
        shown = set()
        for index in matches:
            for nearby in range(max(0, index - 2), min(len(blocks), index + 3)):
                if nearby in shown:
                    continue
                shown.add(nearby)
                label, style, text = blocks[nearby]
                if text:
                    print(f"{label}\t{style}\t{text}")


if __name__ == "__main__":
    main()
