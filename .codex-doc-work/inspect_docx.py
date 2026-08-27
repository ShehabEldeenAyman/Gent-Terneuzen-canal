import sys
from pathlib import Path

from docx import Document


def main() -> None:
    for arg in sys.argv[1:]:
        path = Path(arg)
        doc = Document(path)
        print(f"FILE: {path}")
        print(
            f"PARAGRAPHS {len(doc.paragraphs)} TABLES {len(doc.tables)} "
            f"SECTIONS {len(doc.sections)}"
        )
        for i, paragraph in enumerate(doc.paragraphs):
            text = " ".join(paragraph.text.split())
            if text:
                print(f"P{i:04d}\t{paragraph.style.name}\t{text}")
        for table_i, table in enumerate(doc.tables):
            print(
                f"TABLE {table_i} rows={len(table.rows)} cols={len(table.columns)}"
            )
            for row_i, row in enumerate(table.rows):
                values = [" ".join(cell.text.split()) for cell in row.cells]
                print(f"R{row_i:03d}\t" + " || ".join(values))


if __name__ == "__main__":
    main()
