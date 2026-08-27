from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def normalized(text: str) -> str:
    return " ".join(text.split())


def find_paragraph(doc, exact_text: str) -> Paragraph:
    target = normalized(exact_text)
    for paragraph in doc.paragraphs:
        if normalized(paragraph.text) == target:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def first_run_properties(paragraph: Paragraph):
    for run in paragraph.runs:
        if run.text or run._r.rPr is not None:
            return copy.deepcopy(run._r.rPr) if run._r.rPr is not None else None
    return None


def set_text(paragraph: Paragraph, text: str) -> Paragraph:
    rpr = first_run_properties(paragraph)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def append_text(paragraph: Paragraph, text: str) -> Paragraph:
    rpr = first_run_properties(paragraph)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def insert_after(reference: Paragraph, text: str, donor: Paragraph | None = None) -> Paragraph:
    donor = donor or reference
    new_p = OxmlElement("w:p")
    if donor._p.pPr is not None:
        new_p.append(copy.deepcopy(donor._p.pPr))
    reference._p.addnext(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    run = paragraph.add_run(text)
    rpr = first_run_properties(donor)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def edit_methodology(source: Path, destination: Path) -> None:
    doc = Document(source)
    heading_donor = find_paragraph(doc, "4.5.2 Support Vector Regression")
    body_donor = find_paragraph(
        doc,
        "The SVR model received the same 1,440 flattened values, followed by an "
        "additional feature-wise standardization fitted on the training matrix. A "
        "radial basis function kernel was used with C = 10, epsilon = 0.05 and gamma "
        "= 'scale'. As in XGBoost, 16 independent regressors predicted the scaled "
        "change at each future step. The RBF formulation provides a smooth nonlinear "
        "mapping, in contrast to the piecewise-constant partitions produced by tree "
        "ensembles.",
    )

    methodology_intro = find_paragraph(
        doc,
        "All forecasting techniques were evaluated through a single experimental "
        "contract. Data acquisition, unit normalization, temporal preprocessing, "
        "chronological splitting, window construction, target scaling and evaluation "
        "were therefore performed once and reused unchanged by every model. This design "
        "prevents differences in preprocessing or test-window selection from being "
        "mistaken for differences in model capability. The following subsections first "
        "describe this shared procedure and then report only the elements unique to each "
        "learning technique.",
    )
    set_text(
        methodology_intro,
        "All forecasting techniques were evaluated through a single experimental "
        "contract. Data acquisition, unit normalization, temporal preprocessing, "
        "chronological splitting, window construction and evaluation were therefore "
        "performed once and reused unchanged by every model; model-specific input and "
        "target transformations were fitted within the training data only. This design "
        "prevents differences in preprocessing or test-window selection from being "
        "mistaken for differences in model capability. The following subsections first "
        "describe this shared procedure and then report only the elements unique to each "
        "learning technique.",
    )

    delta_paragraph = find_paragraph(
        doc,
        "The delta-trajectory target was also intended to reduce the tendency of "
        "regression models to revert toward the historical conductivity level and "
        "underestimate variability during rapid transitions. Instead of learning the "
        "absolute level independently at each horizon, every estimator learned the "
        "direction and magnitude of change from the current industrial-site baseline. "
        "The complete-history models could infer local momentum directly from the "
        "preceding target sequence; in the compact 42-hour ablation described in "
        "Section 4.6, this information was summarized explicitly by trailing one-hour "
        "and two-hour target changes. Delta tracking changes the prediction "
        "representation but does not guarantee recovery of every extreme event, so its "
        "benefit was assessed against the same persistence baseline and held-out "
        "trajectories used for all models.",
    )
    set_text(
        delta_paragraph,
        "The delta-trajectory target was also intended to reduce the tendency of "
        "regression models to revert toward the historical conductivity level and "
        "underestimate variability during rapid transitions. Instead of learning the "
        "absolute level independently at each horizon, each task-specific estimator "
        "learned the direction and magnitude of change from the current industrial-site "
        "baseline. The complete-history models could infer local momentum directly from "
        "the preceding target sequence; in the compact 42-hour ablation described in "
        "Section 4.6, this information was summarized explicitly by trailing one-hour "
        "and two-hour target changes. Chronos-2 was evaluated without local weight "
        "training: it forecast target levels in physical units, after which its median "
        "trajectory was expressed as changes from the same current-target baseline for "
        "common evaluation. Delta tracking changes the prediction representation but "
        "does not guarantee recovery of every extreme event, so its benefit was assessed "
        "against the same persistence baseline and held-out trajectories used for all "
        "models.",
    )

    scaling_paragraph = find_paragraph(
        doc,
        "Each sensor was standardized using a mean and standard deviation estimated "
        "only from complete training-period rows. Target changes were scaled separately "
        "at each of the 16 lead times using the corresponding training standard "
        "deviation. Validation and test data did not contribute to either transformation.",
    )
    append_text(
        scaling_paragraph,
        " These external transformations were used by the locally fitted estimators. "
        "Chronos-2 instead received reconstructed physical-scale histories and applied "
        "the pretrained pipeline's internal robust normalization.",
    )

    lag_paragraph = find_paragraph(
        doc,
        "The selected lags were used to align upstream observations for exploratory "
        "visualization and interpretation. They were not imposed on the complete-history "
        "XGBoost, SVR, LSTM or TCN models, which received the original ordered histories "
        "and could learn useful temporal relationships internally. Because the "
        "preliminary lag search was performed before the final chronological evaluation "
        "protocol was fixed, these lag estimates and the associated 42-hour hypothesis "
        "are treated as exploratory. A confirmatory study should estimate and select "
        "lags using training data only and hold them fixed before evaluating the "
        "September test period.",
    )
    set_text(
        lag_paragraph,
        "The selected lags were used to align upstream observations for exploratory "
        "visualization and interpretation. They were not imposed on the complete-history "
        "Ridge, XGBoost, SVR, LSTM, TCN or Chronos-2 configurations, which received the "
        "original ordered histories and could use temporal relationships without a "
        "prescribed shift. Because the preliminary lag search was performed before the "
        "final chronological evaluation protocol was fixed, these lag estimates and the "
        "associated 42-hour hypothesis are treated as exploratory. A confirmatory study "
        "should estimate and select lags using training data only and hold them fixed "
        "before evaluating the September test period.",
    )

    ridge_heading = insert_after(body_donor, "4.5.3 Ridge Regression", heading_donor)
    ridge_body = insert_after(
        ridge_heading,
        "Ridge regression used the same flattened complete-history input as the other "
        "classical models. Feature-wise standardization was fitted on the training "
        "matrix, after which one regularized linear multi-output model predicted the "
        "scaled change at every lead. The L2 penalty was fixed at alpha = 10.0. This "
        "retained all ordered sensor-lag features while providing a linear reference for "
        "the nonlinear and sequence-based architectures.",
        body_donor,
    )

    set_text(
        find_paragraph(doc, "4.5.3 Long Short-Term Memory Network"),
        "4.5.4 Long Short-Term Memory Network",
    )
    set_text(
        find_paragraph(doc, "4.5.4 Temporal Convolutional Network"),
        "4.5.5 Temporal Convolutional Network",
    )
    tcn_body = find_paragraph(
        doc,
        "TCN also preserved the multivariate sequence but replaced recurrence with "
        "causal dilated convolutions. Six residual blocks used dilation factors 1, 2, 4, "
        "8, 16 and 32. Each block contained two causal one-dimensional convolutions with "
        "24 filters and kernel size 5, followed by layer normalization, rectified linear "
        "activation and spatial dropout of 0.10. Residual shortcuts supported gradient "
        "flow. The theoretical receptive field was 505 time steps, exceeding the "
        "available 288-step history; the final encoding could therefore use the complete "
        "72-hour input. The latest encoded time step was passed through a 64-unit dense "
        "layer and a 16-unit output layer. Optimization matched the LSTM objective and "
        "optimizer, with batch size 256, at most 120 epochs, early stopping patience 14 "
        "and learning-rate-reduction patience 5.",
    )
    chronos_heading = insert_after(tcn_body, "4.5.6 Chronos-2 Foundation Model", heading_donor)
    chronos_body_1 = insert_after(
        chronos_heading,
        "Chronos-2 was included as a pretrained probabilistic forecasting reference using "
        "the 120-million-parameter amazon/chronos-2 checkpoint. The industrial-site "
        "history was supplied as the target series, while Terneuzen, Westdorpe, Gent-far "
        "and Gent-near were supplied as past-only covariates. The model received the "
        "unshifted physical-scale histories and generated the complete requested horizon "
        "directly; no future sensor values were supplied.",
        body_donor,
    )
    insert_after(
        chronos_body_1,
        "Inference was performed in float32 with the pretrained weights frozen, so no "
        "canal observations were used for gradient-based fine-tuning. The 0.5 quantile "
        "served as the point forecast, while the 0.1 and 0.9 quantiles were retained for "
        "probabilistic evaluation. In the calibrated configuration, one residual "
        "multiplier per lead was estimated from August only and constrained to the "
        "interval [0, 1], allowing the forecast to shrink toward persistence without "
        "altering the Chronos-2 weights. Both the raw zero-shot and predeclared calibrated "
        "configurations were carried forward to held-out evaluation.",
        body_donor,
    )

    ablation_paragraph = find_paragraph(
        doc,
        "The complete-history XGBoost, SVR, LSTM and TCN models were not supplied with a "
        "manually shifted propagation feature. They could access the entire preceding 72 "
        "hours and determine useful temporal relationships internally. Motivated by the "
        "Terneuzen search interval and the exploratory cross-correlation analysis, a "
        "separate classical-model ablation tested an approximately 42-hour "
        "Terneuzen-to-Ghent delay. Its compact input contained four standardized "
        "features: current industrial-site conductivity, trailing one-hour and two-hour "
        "target changes, and Terneuzen conductivity exactly 42 hours before the forecast "
        "origin. XGBoost and SVR were refitted on this representation using their original "
        "estimator settings and the same training, validation and test targets.",
    )
    set_text(
        ablation_paragraph,
        "The complete-history Ridge, XGBoost, SVR, LSTM, TCN and Chronos-2 configurations "
        "were not supplied with a manually shifted propagation feature. They could access "
        "the entire preceding 72 hours and determine useful temporal relationships without "
        "a prescribed delay. Motivated by the Terneuzen search interval and the exploratory "
        "cross-correlation analysis, a separate ablation tested an approximately 42-hour "
        "Terneuzen-to-Ghent delay. Its compact classical input contained four standardized "
        "features: current industrial-site conductivity, trailing one-hour and two-hour "
        "target changes, and Terneuzen conductivity exactly 42 hours before the forecast "
        "origin. Ridge, XGBoost and SVR were refitted on this representation using their "
        "original estimator settings and the same training, validation and test targets. "
        "For Chronos-2, the matched 42-hour sequence retained the industrial target as the "
        "forecast series and supplied the two target-change histories and shifted "
        "Terneuzen history as past-only covariates, while the pretrained weights remained "
        "frozen.",
    )
    ablation_interpretation = find_paragraph(
        doc,
        "This comparison should be interpreted as a feature-representation experiment "
        "rather than a direct test. The complete 72-hour history already contains the "
        "Terneuzen observation at 42 hours; the ablation changes 1,440 inputs into four "
        "manually selected summaries.",
    )
    set_text(
        ablation_interpretation,
        "This comparison should be interpreted as a feature-representation experiment "
        "rather than a direct test. The complete 72-hour history already contains the "
        "Terneuzen observation at 42 hours. For the classical estimators, the ablation "
        "changes 1,440 inputs into four manually selected summaries; for Chronos-2, the "
        "same four variables are retained as an ordered 42-hour sequence.",
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def edit_architectures(source: Path, destination: Path) -> None:
    doc = Document(source)
    heading_donor = find_paragraph(doc, "5.2. Continuous Function Mapping via Support Vector Regression")
    body_donor = find_paragraph(
        doc,
        "We evaluated Support Vector Regression (SVR) equipped with a Radial Basis "
        "Function (RBF) kernel. Unlike XGBoost, SVR computes a continuous mathematical "
        "function by mapping the input space into a high-dimensional feature space where "
        "a linear regression tube (defined by an insensitive loss function) is fitted.",
    )

    intro = find_paragraph(
        doc,
        "With the raw LDES streams transformed into a physically aligned, delta-tracked "
        "feature matrix, the predictive task shifts to capturing the non-linear fluid "
        "dynamics of the canal. The notebook implements three tabular regressors "
        "(XGBoost, SVR, and MLP) and an exploratory two-stage LSTM "
        "autoencoder/forecaster. Furthermore, to ensure the reproducibility of these "
        "models within the Water Data Space, the pipeline can subsequently be documented "
        "using semantic provenance vocabularies.",
    )
    set_text(
        intro,
        "With the raw LDES streams transformed into a physically aligned, delta-tracked "
        "feature matrix, the predictive task shifts to representing both linear and "
        "nonlinear temporal relationships in the canal. The modelling suite combines "
        "Ridge regression, XGBoost and SVR with the LSTM and TCN sequence architectures, "
        "and includes Chronos-2 as a frozen zero-shot foundation-model benchmark. To "
        "support reproducibility within the Water Data Space, the resulting local and "
        "externally pretrained configurations can be documented using semantic "
        "provenance vocabularies.",
    )

    tree_last = find_paragraph(
        doc,
        "Because the forecast is anchored to the real-time baseline value Y(t) measured "
        "at the target site, delta reconstruction can reduce absolute-scale boundary "
        "effects. Whether it improves extreme-event prediction must nevertheless be "
        "established from held-out evaluation results.",
    )
    ridge_heading = insert_after(
        tree_last,
        "5.2. Regularized Linear Forecasting via Ridge Regression",
        heading_donor,
    )
    ridge_body_1 = insert_after(
        ridge_heading,
        "Ridge regression was included as a regularized linear reference for the "
        "high-dimensional and strongly correlated lag representation. After "
        "training-only feature standardization, it estimates a coefficient matrix for "
        "the complete multi-step delta trajectory by minimizing squared prediction error "
        "together with an L2 penalty on the coefficients. The regularization strength was "
        "fixed at alpha = 10.0.",
        body_donor,
    )
    insert_after(
        ridge_body_1,
        "Unlike feature selection, L2 regularization retains every sensor-lag input but "
        "shrinks unstable coefficients toward zero. Ridge therefore provides an explicit "
        "linear benchmark against which the additional functional flexibility of kernels, "
        "trees, recurrent networks, convolutions and the pretrained foundation model can "
        "be evaluated under the same chronological protocol.",
        body_donor,
    )

    set_text(
        find_paragraph(doc, "5.2. Continuous Function Mapping via Support Vector Regression"),
        "5.3. Continuous Function Mapping via Support Vector Regression",
    )
    set_text(
        find_paragraph(doc, "5.3. Temporal Convolutional Network"),
        "5.4. Temporal Convolutional Network",
    )
    set_text(
        find_paragraph(doc, "5.4. LSTM Autoencoder and Forecaster"),
        "5.5. LSTM Autoencoder and Forecaster",
    )

    lstm_last = find_paragraph(
        doc,
        "Importantly, the LSTM was not supplied with a manually shifted 42-hour "
        "Terneuzen feature. The network received the unshifted 72-hour history of every "
        "station and was allowed to regulate its temporal memory through its recurrent "
        "input, forget and output gates. The model could therefore learn statistical "
        "dependencies at approximately 42 hours if supported by the data, but that delay "
        "was neither prescribed nor guaranteed by the architecture. Consequently, the "
        "experiment evaluates whether useful propagation information can emerge from the "
        "complete observation history without embedding the earlier cross-correlation "
        "result directly into the model.",
    )
    chronos_heading = insert_after(
        lstm_last,
        "5.6. Zero-Shot Probabilistic Forecasting with Chronos-2",
        heading_donor,
    )
    chronos_body_1 = insert_after(
        chronos_heading,
        "Chronos-2 was included as a pretrained time-series foundation model rather than "
        "a locally optimized network. The 120-million-parameter amazon/chronos-2 "
        "checkpoint first applies robust per-series scaling, divides each series into "
        "non-overlapping temporal patches and maps those patches to learned embeddings. "
        "Its transformer blocks alternate attention over time within a series with group "
        "attention across related series, enabling target and covariate histories to be "
        "processed within one forecasting task.",
        body_donor,
    )
    chronos_body_2 = insert_after(
        chronos_body_1,
        "For this study, industrial-site conductivity was the target series and the four "
        "upstream conductivity histories were past-only covariates. The complete unshifted "
        "history was passed in physical units, with no future covariate values and no "
        "manually prescribed propagation lag. The frozen model generated direct quantile "
        "forecasts over the complete requested horizon. The median (0.5 quantile) was used "
        "as the point forecast, while the 0.1 and 0.9 quantiles were retained to describe "
        "forecast uncertainty.",
        body_donor,
    )
    insert_after(
        chronos_body_2,
        "No canal observations were used to update the Chronos-2 parameters. A separate, "
        "predeclared August-only calibration estimated one residual multiplier per lead, "
        "constrained to [0, 1], so that calibrated residuals could only shrink toward the "
        "persistence forecast. The raw zero-shot output was retained alongside this "
        "calibrated configuration for held-out evaluation; neither configuration changes "
        "the pretrained model weights.",
        body_donor,
    )

    provenance_heading = find_paragraph(doc, "5.5. Semantic Model Provenance with MLDCAT-AP")
    set_text(provenance_heading, "5.7. Semantic Model Provenance with MLDCAT-AP")
    provenance_intro = find_paragraph(
        doc,
        "A critical challenge in modern data science is the isolation of machine learning "
        "artifacts from their underlying data architecture. To integrate the LSTM "
        "artifacts into the wider data space, we provide an MLDCAT-AP 2.0 Turtle. This "
        "record describes the published Seq2Seq forecaster",
    )
    set_text(
        provenance_intro,
        "A critical challenge in modern data science is the isolation of machine learning "
        "artifacts from their underlying data architecture. To integrate the locally "
        "trained artifacts and the Chronos-2 inference configuration into the wider data "
        "space, we provide MLDCAT-AP 2.0 Turtle metadata. The existing record describes "
        "the published Seq2Seq forecaster, while model-specific records can distinguish "
        "trained artifacts from externally pretrained dependencies.",
    )
    provenance_repository = find_paragraph(
        doc,
        "To support transparent provenance and open science principles, the metadata "
        "explicitly defines the creator's identity, linking the model's authorship to "
        "institutional affiliations. Furthermore, the graph captures the artifact's "
        "CC-BY 4.0 open license, semantic versioning, cryptographic checksums, and its "
        "source code repository.",
    )
    insert_after(
        provenance_repository,
        "For Ridge, the provenance record should include the training split, feature and "
        "target scalers, alpha value and software version. For Chronos-2, it should record "
        "the upstream model identifier and revision, package version, license and "
        "checkpoint checksum, zero-shot status, input and covariate schema, requested "
        "quantiles, inference device and numerical precision, and the separate "
        "validation-only calibration parameters. These fields make clear which components "
        "were learned locally and which were inherited from the pretrained checkpoint.",
        body_donor,
    )
    provenance_boundaries = find_paragraph(
        doc,
        "Crucially, the MLDCAT-AP record formally encodes both the intended use and the "
        "operational boundaries of the deployed model to ensure safe adoption. The "
        "metadata specifies the model's input window to produce a 288-step (72-hour) "
        "forecast horizon at a 15-minute resolution. It also explicitly asserts the "
        "model's limitations such as forecasting accuracy degradation beyond the "
        "24-hour mark and mandates the use of the original Minmax scaler for correct "
        "input-output transformation. These limitations are recorded explicitly, "
        "ensuring that downstream consumer applications and researchers can query and "
        "evaluate the model's constraints prior to operational use.",
    )
    set_text(
        provenance_boundaries,
        "Crucially, the MLDCAT-AP record formally encodes the intended use and operational "
        "boundaries of each model configuration. The metadata specifies the input window, "
        "forecast horizon, temporal resolution, required preprocessing and scalers, and "
        "the target and covariate schema. Until held-out results are finalized, empirical "
        "accuracy limitations and comparative performance remain unspecified; supported "
        "limitations can be added to the record after evaluation without changing the "
        "method definition.",
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> None:
    if len(sys.argv) != 5:
        raise SystemExit(
            "Usage: edit_models.py methodology_in architectures_in methodology_out architectures_out"
        )
    methodology_in, architectures_in, methodology_out, architectures_out = map(
        Path, sys.argv[1:]
    )
    edit_methodology(methodology_in, methodology_out)
    edit_architectures(architectures_in, architectures_out)


if __name__ == "__main__":
    main()
